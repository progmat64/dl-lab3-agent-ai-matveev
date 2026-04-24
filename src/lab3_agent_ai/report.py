from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


def _add_table_from_df(doc: Document, df: pd.DataFrame) -> None:
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    for i, col in enumerate(df.columns):
        table.cell(0, i).text = str(col)
    for _, row_data in df.iterrows():
        row = table.add_row().cells
        for i, value in enumerate(row_data):
            row[i].text = str(value)


def build_report(df: pd.DataFrame, summary: pd.DataFrame, figure_paths: list[Path], output_path: Path) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(1.8)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    title = doc.add_heading("Отчет по лабораторной работе 3: Agent AI", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("1. Цель эксперимента", level=1)
    doc.add_paragraph(
        "Цель работы - сравнить baseline-подход и агентный подход при подготовке "
        "научно-аналитического обзора, а также проверить эффект evaluator."
    )

    doc.add_heading("2. Условия запуска", level=1)
    model_names = ", ".join(sorted(str(model) for model in df["model"].dropna().unique()))
    doc.add_paragraph(f"Используемая модель: {model_names}.")
    doc.add_paragraph(f"Количество тем: {df['topic'].nunique()}.")
    doc.add_paragraph(
        "Оцениваемые конфигурации: "
        + ", ".join(sorted(str(mode) for mode in df["mode"].dropna().unique()))
        + "."
    )

    doc.add_heading("3. Конфигурации", level=1)
    for item in [
        "baseline: один ответ LLM на основе общего контекста Wikipedia.",
        "agent: пошаговый поиск, состояние, notes и JSON trace.",
        "agent + evaluator: agent с внутренней проверкой качества и возможной доработкой ответа.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("4. Сводные результаты", level=1)
    _add_table_from_df(doc, summary.reset_index())

    doc.add_heading("5. Результаты по запускам", level=1)
    display_cols = [
        "topic",
        "mode",
        "correctness",
        "groundedness",
        "completeness",
        "coverage_of_required_fields",
        "rubric",
        "n_steps",
        "latency",
        "tool_errors",
    ]
    _add_table_from_df(doc, df[display_cols].round(3))

    doc.add_heading("6. Графики", level=1)
    for fig in figure_paths:
        if fig.exists():
            doc.add_paragraph(fig.name)
            doc.add_picture(str(fig), width=Cm(15))

    doc.add_heading("7. Вывод", level=1)
    best_rubric_mode = summary["rubric"].idxmax()
    best_grounded_mode = summary["groundedness"].idxmax()
    fastest_mode = summary["latency"].idxmin()
    most_compact_mode = summary["n_steps"].idxmin()
    baseline = summary.loc["baseline"]
    agent = summary.loc["agent"]
    agent_eval = summary.loc["agent_evaluator"]

    doc.add_paragraph(
        f"По средней итоговой оценке rubric лучший результат показала конфигурация "
        f"{best_rubric_mode}: {summary.loc[best_rubric_mode, 'rubric']:.3f}. "
        f"Baseline получил {baseline['rubric']:.3f}, agent - {agent['rubric']:.3f}, "
        f"agent + evaluator - {agent_eval['rubric']:.3f}."
    )
    doc.add_paragraph(
        f"Главное отличие agent-режима от baseline - качество grounding и source consistency. "
        f"Groundedness вырос с {baseline['groundedness']:.3f} у baseline до "
        f"{agent['groundedness']:.3f} у agent, а source consistency - с "
        f"{baseline['source_consistency']:.3f} до {agent['source_consistency']:.3f}. "
        "Это связано с тем, что agent явно передает в prompt найденные OpenAlex-источники и "
        "сохраняет промежуточные notes."
    )
    doc.add_paragraph(
        f"Agent + evaluator оказался полезен как контроль качества, но в данном запуске не улучшил "
        f"средний rubric относительно agent без evaluator: {agent_eval['rubric']:.3f} против "
        f"{agent['rubric']:.3f}. При этом он добавил один шаг trace и увеличил среднюю latency "
        f"с {agent['latency']:.3f} с до {agent_eval['latency']:.3f} с."
    )
    doc.add_paragraph(
        f"Самым быстрым режимом оказался {fastest_mode}, а самым компактным по числу шагов - "
        f"{most_compact_mode}. Для итоговой лабораторной интерпретации важно разделять качество "
        "итогового текста и качество процесса: baseline может давать связный обзор, но без "
        "источников проигрывает по проверяемости; agent дороже по шагам, но лучше обосновывает выводы."
    )

    doc.save(output_path)
    return output_path
